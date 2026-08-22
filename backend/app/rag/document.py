"""
Mode 2: Session-scoped FAISS document store for uploaded documents.
Each user session gets its own in-memory FAISS index that is destroyed on clear.
"""
import io
import logging
from typing import Optional
from dataclasses import dataclass, field

from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.rag.embeddings import LocalEmbeddings

logger = logging.getLogger(__name__)


@dataclass
class SessionDocument:
    """Holds the FAISS index and metadata for an uploaded document."""
    filename: str
    page_count: int
    chunk_count: int
    vectorstore: FAISS
    chunks: list[Document] = field(default_factory=list)


class SessionDocumentStore:
    """
    Manages per-session FAISS indexes for uploaded documents.
    Thread-safe for single-process deployments (typical for dev/demo).
    """

    def __init__(self):
        self._sessions: dict[str, SessionDocument] = {}
        self._text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        self._embeddings = LocalEmbeddings()

    def _extract_pdf(self, file_bytes: bytes) -> tuple[str, int]:
        """Extract text from a PDF. Returns (full_text, page_count)."""
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages), len(reader.pages)

    def _extract_docx(self, file_bytes: bytes) -> tuple[str, int]:
        """Extract text from a DOCX. Returns (full_text, page_count_estimate)."""
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        # DOCX doesn't have a reliable page count; estimate from char length
        page_estimate = max(1, len(full_text) // 3000)
        return full_text, page_estimate

    def add_document(self, session_id: str, file_bytes: bytes, filename: str) -> SessionDocument:
        """
        Parse, chunk, embed, and index a document for the given session.
        Replaces any previously uploaded document for this session.
        """
        # Remove existing document if any
        self.remove_document(session_id)

        # Extract text based on file type
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        if ext == "pdf":
            full_text, page_count = self._extract_pdf(file_bytes)
        elif ext in ("docx", "doc"):
            full_text, page_count = self._extract_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: .{ext}. Use PDF or DOCX.")

        if not full_text.strip():
            raise ValueError("Could not extract any text from the uploaded document.")

        # Chunk the text
        chunks = self._text_splitter.create_documents(
            texts=[full_text],
            metadatas=[{"source": filename}],
        )

        logger.info(f"Session {session_id}: chunked '{filename}' into {len(chunks)} chunks")

        # Build FAISS index
        vectorstore = FAISS.from_documents(chunks, self._embeddings)

        session_doc = SessionDocument(
            filename=filename,
            page_count=page_count,
            chunk_count=len(chunks),
            vectorstore=vectorstore,
            chunks=chunks,
        )
        self._sessions[session_id] = session_doc
        return session_doc

    def query(self, session_id: str, query_text: str, k: int = 4) -> list[Document]:
        """Retrieve relevant chunks from the session's document."""
        session_doc = self._sessions.get(session_id)
        if session_doc is None:
            return []
        return session_doc.vectorstore.similarity_search(query_text, k=k)

    def has_document(self, session_id: str) -> bool:
        """Check if a session has an active document."""
        return session_id in self._sessions

    def get_document_info(self, session_id: str) -> Optional[dict]:
        """Get metadata about the active document for a session."""
        session_doc = self._sessions.get(session_id)
        if session_doc is None:
            return None
        return {
            "filename": session_doc.filename,
            "page_count": session_doc.page_count,
            "chunk_count": session_doc.chunk_count,
        }

    def remove_document(self, session_id: str) -> bool:
        """Remove the document and FAISS index for a session."""
        if session_id in self._sessions:
            del self._sessions[session_id]
            logger.info(f"Session {session_id}: document cleared")
            return True
        return False


# Global singleton — one per process
document_store = SessionDocumentStore()
